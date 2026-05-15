"""
Minimal document parser for the Document-first TamperShield pipeline.

This module converts PDF, DOCX, and image inputs into shared document models.
It does not implement page alignment, content comparison, OCR, table matching,
cell comparison, evidence indexing, or report generation.
"""

from pathlib import Path
import contextlib
import io
import subprocess
import tempfile
from typing import Any

import fitz
from docx import Document

from core.document_models import (
    BoundingBox,
    DocumentElement,
    DocumentPage,
    ElementType,
    FileType,
    ParsedDocument,
)


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}


def detect_file_type(file_path: str | Path) -> FileType:
    """Detect a supported document type from the filename extension only."""
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    return "unknown"


def parse_document(
    file_path: str | Path,
    enable_ocr: bool = False,
    enable_preprocess: bool = False,
) -> ParsedDocument:
    """Parse a file into a ParsedDocument using the minimal Phase 3 entrypoint."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(str(path))

    file_type = detect_file_type(path)

    if file_type == "pdf":
        return parse_pdf_document(path)
    if file_type == "docx":
        return parse_docx_document(path)
    if file_type == "image":
        return parse_image_document(
            path,
            enable_ocr=enable_ocr,
            enable_preprocess=enable_preprocess,
        )

    return ParsedDocument(
        file_path=str(path),
        file_type="unknown",
        metadata={
            "source_file": str(path),
            "source_format": "unknown",
        },
    )


def parse_pdf_document(file_path: str | Path) -> ParsedDocument:
    """Parse PDF pages and text blocks with PyMuPDF."""
    path = Path(file_path)
    pages: list[DocumentPage] = []

    with fitz.open(str(path)) as pdf:
        for page_index, pdf_page in enumerate(pdf):
            page_number = page_index + 1
            plain_text = pdf_page.get_text("text") or ""
            document_page = DocumentPage(
                page_number=page_number,
                plain_text=plain_text,
                metadata={
                    "source_file": str(path),
                    "source_format": "pdf",
                    "page_index": page_index,
                },
            )

            blocks = pdf_page.get_text("blocks") or []
            text_block_count = 0
            for block_index, block in enumerate(blocks):
                block_text = _pdf_block_text(block)
                if not block_text.strip():
                    continue

                bbox = BoundingBox(
                    page_number=page_number,
                    x0=float(block[0]),
                    y0=float(block[1]),
                    x1=float(block[2]),
                    y1=float(block[3]),
                )
                document_page.add_element(
                    DocumentElement(
                        element_id=f"pdf-page-{page_number}-block-{block_index}",
                        element_type="paragraph",
                        page_number=page_number,
                        text=block_text,
                        bbox=bbox,
                        raw=block,
                        metadata={
                            "source_file": str(path),
                            "source_format": "pdf",
                            "page_index": page_index,
                            "block_index": block_index,
                        },
                    )
                )
                text_block_count += 1

            for image_index, image_element in enumerate(
                _extract_pdf_image_elements(
                    pdf_page=pdf_page,
                    path=path,
                    page_number=page_number,
                    page_index=page_index,
                )
            ):
                image_element.element_id = f"pdf-page-{page_number}-image-{image_index}"
                document_page.add_element(image_element)

            for table_index, table_element in enumerate(
                _extract_pdf_table_elements(
                    pdf_page=pdf_page,
                    path=path,
                    page_number=page_number,
                    page_index=page_index,
                )
            ):
                table_element.element_id = f"pdf-page-{page_number}-table-{table_index}"
                document_page.add_element(table_element)

            if text_block_count == 0:
                document_page.add_element(
                    DocumentElement(
                        element_id=f"pdf-page-{page_number}-text-0",
                        element_type="paragraph" if plain_text.strip() else "unknown",
                        page_number=page_number,
                        text=plain_text,
                        metadata={
                            "source_file": str(path),
                            "source_format": "pdf",
                            "page_index": page_index,
                        },
                    )
                )

            pages.append(document_page)

    return ParsedDocument(
        file_path=str(path),
        file_type="pdf",
        pages=pages,
        metadata={
            "source_file": str(path),
            "source_format": "pdf",
        },
    )


def parse_docx_document(file_path: str | Path) -> ParsedDocument:
    """Parse DOCX with rendered-PDF pagination when available.

    Falls back to native single-page parsing when rendering is unavailable.
    """
    path = Path(file_path)

    try:
        return _parse_docx_via_rendered_pdf(path)
    except Exception as exc:
        parsed = _parse_docx_native_single_page(path)
        parsed.metadata["pagination"] = "native_single_page_fallback"
        parsed.metadata["pagination_fallback_reason"] = type(exc).__name__
        parsed.metadata["pagination_fallback_message"] = str(exc)
        return parsed


def _parse_docx_native_single_page(file_path: str | Path) -> ParsedDocument:
    """Parse DOCX native paragraphs and tables into a single logical page."""
    path = Path(file_path)
    doc = Document(str(path))
    page = DocumentPage(
        page_number=1,
        metadata={
            "source_file": str(path),
            "source_format": "docx",
            "page_index": 0,
            "pagination": "not_available_native_docx",
        },
    )

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        page.add_element(
            DocumentElement(
                element_id=f"docx-paragraph-{paragraph_index}",
                element_type="paragraph",
                page_number=1,
                text=text,
                metadata={
                    "source_file": str(path),
                    "source_format": "docx",
                    "pagination": "not_available_native_docx",
                    "paragraph_index": paragraph_index,
                },
            )
        )
    # Phase 3 minimal implementation: paragraphs and tables are extracted in
    # separate passes. Original body-order traversal can be added later.
    for table_index, table in enumerate(doc.tables):
        page.add_element(
            _docx_table_to_element(
                table=table,
                path=path,
                table_index=table_index,
                page_number=1,
                page_index=0,
                pagination="not_available_native_docx",
                page_index_estimated=False,
            )
        )

    for image_index, image_element in enumerate(
        _extract_docx_image_elements(
            doc=doc,
            path=path,
            page_number=1,
            page_index=0,
            pagination="not_available_native_docx",
            page_index_estimated=False,
        )
    ):
        image_element.element_id = f"docx-image-{image_index}"
        page.add_element(image_element)

    page.plain_text = "\n".join(
        element.text for element in page.elements if element.text.strip()
    )

    return ParsedDocument(
        file_path=str(path),
        file_type="docx",
        pages=[page],
        metadata={
            "source_file": str(path),
            "source_format": "docx",
            "pagination": "not_available_native_docx",
        },
    )


def _convert_docx_to_pdf_with_libreoffice(
    docx_path: Path,
    output_dir: Path,
) -> Path:
    """Convert a DOCX file to PDF in a temporary directory with LibreOffice."""
    errors: list[str] = []
    for executable in ("soffice", "libreoffice"):
        command = [
            executable,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=False,
            )
        except FileNotFoundError as exc:
            errors.append(f"{executable}: {exc}")
            continue

        expected_pdf = output_dir / f"{docx_path.stem}.pdf"
        if result.returncode == 0 and expected_pdf.exists():
            return expected_pdf

        generated_pdfs = sorted(output_dir.glob("*.pdf"))
        if result.returncode == 0 and generated_pdfs:
            return generated_pdfs[0]

        errors.append(
            (
                f"{executable}: returncode={result.returncode}; "
                f"stdout={result.stdout.strip()}; "
                f"stderr={result.stderr.strip()}"
            )
        )

    raise RuntimeError("DOCX to PDF conversion failed: " + " | ".join(errors))


def _parse_docx_via_rendered_pdf(file_path: str | Path) -> ParsedDocument:
    """Render DOCX to a temporary PDF and parse PDF-like pages."""
    original_path = Path(file_path)
    with tempfile.TemporaryDirectory() as temporary_dir:
        output_dir = Path(temporary_dir)
        rendered_pdf_path = _convert_docx_to_pdf_with_libreoffice(
            original_path,
            output_dir,
        )
        rendered_pdf_doc = parse_pdf_document(rendered_pdf_path)

        metadata = {
            "source_file": str(original_path),
            "source_format": "docx",
            "pagination": "rendered_pdf",
            "render_backend": "libreoffice",
            "rendered_pdf_page_count": len(rendered_pdf_doc.pages),
        }
        native_elements = _extract_docx_native_supplement_elements(
            original_path,
            pagination="rendered_pdf",
        )
        if native_elements and rendered_pdf_doc.pages:
            rendered_pdf_doc.pages[0].elements.extend(native_elements)

        for page in rendered_pdf_doc.pages:
            page.metadata.update(metadata)
            for element in page.elements:
                element.metadata.update(metadata)

        metadata["native_docx_table_count"] = sum(
            1 for element in native_elements if element.element_type == "table"
        )
        metadata["native_docx_image_count"] = sum(
            1 for element in native_elements if element.element_type == "image"
        )

        return ParsedDocument(
            file_path=str(original_path),
            file_type="docx",
            pages=rendered_pdf_doc.pages,
            metadata=metadata,
        )


def parse_image_document(
    file_path: str | Path,
    enable_ocr: bool = False,
    ocr_lang: str = "ch",
    ocr_use_gpu: bool = False,
    enable_preprocess: bool = False,
    preprocess_mode: str = "gray",
    enable_deskew: bool = True,
    preserve_ocr_blocks: bool = True,
) -> ParsedDocument:
    """Parse an image document, optionally with OCR-backed elements."""
    path = Path(file_path)

    if not enable_ocr:
        return _parse_image_document_without_ocr(path)

    return _parse_image_document_with_ocr(
        path=path,
        ocr_lang=ocr_lang,
        ocr_use_gpu=ocr_use_gpu,
        enable_preprocess=enable_preprocess,
        preprocess_mode=preprocess_mode,
        enable_deskew=enable_deskew,
        preserve_ocr_blocks=preserve_ocr_blocks,
    )


def _parse_image_document_without_ocr(path: Path) -> ParsedDocument:
    """Register an image as a single-page document without running OCR."""
    page = DocumentPage(
        page_number=1,
        page_image_path=str(path),
        metadata={
            "source_file": str(path),
            "source_format": "image",
            "page_index": 0,
        },
    )
    page.add_element(
        DocumentElement(
            element_id="image-page-1-image-0",
            element_type="image",
            page_number=1,
            metadata={
                "source_file": str(path),
                "source_format": "image",
                "page_index": 0,
            },
        )
    )

    return ParsedDocument(
        file_path=str(path),
        file_type="image",
        pages=[page],
        metadata={
            "source_file": str(path),
            "source_format": "image",
        },
    )


def _parse_image_document_with_ocr(
    path: Path,
    ocr_lang: str = "ch",
    ocr_use_gpu: bool = False,
    enable_preprocess: bool = False,
    preprocess_mode: str = "gray",
    enable_deskew: bool = True,
    preserve_ocr_blocks: bool = True,
) -> ParsedDocument:
    """Parse an image as a single-page OCR-backed ParsedDocument."""
    from core.ocr_engine import (
        build_pp_structure,
        extract_tables_with_metadata,
        parse_layout_to_blocks,
    )

    engine = build_pp_structure(lang=ocr_lang, use_gpu=ocr_use_gpu)
    blocks = parse_layout_to_blocks(engine, str(path))

    preprocess_metadata = {
        "preprocess_enabled": enable_preprocess,
        "preprocess_mode": preprocess_mode,
        "deskew_enabled": enable_deskew,
        "preprocess_applied": False,
        "preprocess_output_in_memory": False,
        "ocr_input": "original_image_path",
        "preprocess_note": (
            "Phase 14c records preprocessing parameters only; "
            "OCR used original image path to avoid intermediate file writes."
        ),
    }

    page = DocumentPage(
        page_number=1,
        page_image_path=str(path),
        metadata={
            "source_file": str(path),
            "source_format": "image",
            "page_index": 0,
            "ocr_enabled": True,
            "ocr_page_index": 0,
            "native_text_available": False,
            "ocr_applied": True,
            "ocr_backend": "paddleocr_ppstructure",
            "ocr_lang": ocr_lang,
            "ocr_use_gpu": ocr_use_gpu,
            "preprocess": dict(preprocess_metadata),
        },
    )

    block_texts: list[str] = []
    for block_index, block in enumerate(blocks):
        block_text = _ocr_block_text(block)
        if block_text.strip():
            block_texts.append(block_text.strip())

        block_type = _ocr_block_type(block)
        bbox = _ocr_block_bbox(block, page_number=1)
        page.add_element(
            DocumentElement(
                element_id=f"ocr-page-1-block-{block_index}",
                element_type=_ocr_block_to_element_type(block_type),
                page_number=1,
                text=block_text,
                bbox=bbox,
                raw=block if preserve_ocr_blocks else None,
                metadata={
                    "source_file": str(path),
                    "source_format": "image",
                    "page_index": 0,
                    "ocr_enabled": True,
                    "ocr_block_index": block_index,
                    "ocr_block_type": block_type,
                    "ocr_confidence": _ocr_block_confidence(block),
                    "bbox_source": _ocr_block_bbox_source(block),
                    "preprocess": dict(preprocess_metadata),
                },
            )
        )

    tables = extract_tables_with_metadata(blocks)
    for table_index, table_item in enumerate(tables):
        table_text = _flatten_ocr_table_text(table_item)
        if table_text.strip():
            block_texts.append(table_text.strip())

        page.add_element(
            DocumentElement(
                element_id=f"ocr-page-1-table-{table_index}",
                element_type="table",
                page_number=1,
                text=table_text,
                bbox=_ocr_table_bbox(table_item, page_number=1),
                raw={
                    "df": table_item.get("df"),
                    "html": table_item.get("html"),
                    "spans": table_item.get("spans"),
                },
                metadata={
                    "source_file": str(path),
                    "source_format": "image",
                    "page_index": 0,
                    "ocr_enabled": True,
                    "table_index": table_index,
                    "html": table_item.get("html"),
                    "spans": table_item.get("spans"),
                    "cell_box_list": table_item.get("cell_box_list"),
                    "table_ocr_pred": table_item.get("table_ocr_pred"),
                    "source_style": table_item.get("source_style"),
                    "preprocess": dict(preprocess_metadata),
                },
            )
        )

    if not page.elements:
        page.add_element(
            DocumentElement(
                element_id="ocr-page-1-image-0",
                element_type="image",
                page_number=1,
                metadata={
                    "source_file": str(path),
                    "source_format": "image",
                    "page_index": 0,
                    "ocr_enabled": True,
                    "ocr_note": "No OCR blocks were returned.",
                    "preprocess": dict(preprocess_metadata),
                },
            )
        )

    page.plain_text = "\n".join(block_texts)

    return ParsedDocument(
        file_path=str(path),
        file_type="image",
        pages=[page],
        metadata={
            "source_file": str(path),
            "source_format": "image",
            "ocr_enabled": True,
            "ocr_lang": ocr_lang,
            "ocr_use_gpu": ocr_use_gpu,
            "preprocess_enabled": enable_preprocess,
            "preprocess_mode": preprocess_mode,
            "deskew_enabled": enable_deskew,
            "preserve_ocr_blocks": preserve_ocr_blocks,
            "ocr_backend": "paddleocr_ppstructure",
            "preprocess": dict(preprocess_metadata),
        },
    )


def _ocr_get(obj: Any, key: str, default: Any = None) -> Any:
    if obj is None:
        return default

    if isinstance(obj, dict):
        return obj.get(key, default)

    if hasattr(obj, key):
        return getattr(obj, key)

    try:
        return obj[key]
    except Exception:
        return default


def _ocr_block_text(block: Any) -> str:
    text = _ocr_get(block, "text", None)
    if text is not None:
        return _coerce_ocr_text(text)

    res = _ocr_get(block, "res", None)
    if res is not None:
        res_text = _ocr_get(res, "text", None)
        if res_text is not None:
            return _coerce_ocr_text(res_text)

        rec_texts = _ocr_get(res, "rec_texts", None)
        if rec_texts is not None:
            return _coerce_ocr_text(rec_texts)

    rec_texts = _ocr_get(block, "rec_texts", None)
    if rec_texts is not None:
        return _coerce_ocr_text(rec_texts)

    return ""


def _coerce_ocr_text(value: Any) -> str:
    if value is None:
        return ""

    if isinstance(value, str):
        return value

    if isinstance(value, (list, tuple)):
        parts = [str(item).strip() for item in value if str(item).strip()]
        return " ".join(parts)

    return str(value)


def _ocr_block_bbox(block: Any, page_number: int) -> BoundingBox | None:
    bbox_value = _first_ocr_field(
        block,
        keys=("bbox", "box", "layout_bbox", "coordinate"),
    )
    return _bbox_from_value(bbox_value, page_number=page_number)


def _ocr_block_bbox_source(block: Any) -> str:
    for key in ("bbox", "box", "layout_bbox", "coordinate"):
        if _first_ocr_field(block, keys=(key,)) is not None:
            return key
    return ""


def _ocr_block_type(block: Any) -> str:
    block_type = _first_ocr_field(block, keys=("type", "label"))
    if block_type is None:
        return ""
    return str(block_type).strip().lower()


def _ocr_block_confidence(block: Any) -> Any:
    return _first_ocr_field(
        block,
        keys=("confidence", "score", "prob", "rec_scores"),
    )


def _first_ocr_field(obj: Any, keys: tuple[str, ...]) -> Any:
    for key in keys:
        value = _ocr_get(obj, key, None)
        if value is not None:
            return value

    res = _ocr_get(obj, "res", None)
    if res is not None:
        for key in keys:
            value = _ocr_get(res, key, None)
            if value is not None:
                return value

    return None


def _ocr_block_to_element_type(block_type: str) -> ElementType:
    normalized = block_type.strip().lower()

    if normalized == "title":
        return "title"
    if normalized in {"text", "paragraph"}:
        return "paragraph"
    if normalized == "table":
        return "table"
    if normalized in {"figure", "image"}:
        return "image"
    if normalized in {"seal", "signature"}:
        return "signature"
    if normalized == "header":
        return "header"
    if normalized == "footer":
        return "footer"

    return "unknown"


def _ocr_table_bbox(table_item: dict[str, Any], page_number: int) -> BoundingBox | None:
    return _bbox_from_value(table_item.get("bbox"), page_number=page_number)


def _bbox_from_value(value: Any, page_number: int) -> BoundingBox | None:
    if value is None:
        return None

    if isinstance(value, (list, tuple)) and len(value) == 4:
        try:
            x0, y0, x1, y1 = [float(item) for item in value]
        except (TypeError, ValueError):
            return None
        return BoundingBox(
            page_number=page_number,
            x0=x0,
            y0=y0,
            x1=x1,
            y1=y1,
        )

    return None


def _extract_pdf_image_elements(
    pdf_page: Any,
    path: Path,
    page_number: int,
    page_index: int,
) -> list[DocumentElement]:
    elements: list[DocumentElement] = []
    page_dict = pdf_page.get_text("dict") or {}
    blocks = page_dict.get("blocks", [])

    for image_index, block in enumerate(blocks):
        if block.get("type") != 1:
            continue

        bbox = _bbox_from_value(block.get("bbox"), page_number=page_number)
        elements.append(
            DocumentElement(
                element_id=f"pdf-page-{page_number}-image-{image_index}",
                element_type="image",
                page_number=page_number,
                bbox=bbox,
                raw=None,
                metadata={
                    "source_file": str(path),
                    "source_format": "pdf",
                    "page_index": page_index,
                    "image_index": image_index,
                    "width": block.get("width"),
                    "height": block.get("height"),
                    "extension": block.get("ext"),
                    "xref": block.get("xref"),
                    "detection_method": "pymupdf",
                },
            )
        )

    return elements


def _extract_pdf_table_elements(
    pdf_page: Any,
    path: Path,
    page_number: int,
    page_index: int,
) -> list[DocumentElement]:
    elements = _extract_pdf_tables_with_pymupdf(
        pdf_page=pdf_page,
        path=path,
        page_number=page_number,
        page_index=page_index,
    )
    if elements:
        return elements

    return _extract_pdf_table_like_blocks(
        pdf_page=pdf_page,
        path=path,
        page_number=page_number,
        page_index=page_index,
    )


def _extract_pdf_tables_with_pymupdf(
    pdf_page: Any,
    path: Path,
    page_number: int,
    page_index: int,
) -> list[DocumentElement]:
    if not hasattr(pdf_page, "find_tables"):
        return []

    try:
        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            table_finder = pdf_page.find_tables()
    except Exception:
        return []

    tables = getattr(table_finder, "tables", None) or []
    elements: list[DocumentElement] = []
    for table_index, table in enumerate(tables):
        rows = _extract_pymupdf_table_rows(table)
        text = _flatten_table_text(rows)
        metadata = _table_metadata(
            rows=rows,
            detection_method="pymupdf",
            page_index_estimated=False,
        )
        metadata.update(
            {
                "source_file": str(path),
                "source_format": "pdf",
                "page_index": page_index,
                "table_index": table_index,
            }
        )
        elements.append(
            DocumentElement(
                element_id=f"pdf-page-{page_number}-table-{table_index}",
                element_type="table",
                page_number=page_number,
                text=text,
                bbox=_bbox_from_value(getattr(table, "bbox", None), page_number=page_number),
                raw=rows,
                metadata=metadata,
            )
        )

    return elements


def _extract_pymupdf_table_rows(table: Any) -> list[list[str]]:
    try:
        rows = table.extract()
    except Exception:
        rows = []

    return _normalize_table_rows(rows)


def _extract_pdf_table_like_blocks(
    pdf_page: Any,
    path: Path,
    page_number: int,
    page_index: int,
) -> list[DocumentElement]:
    elements: list[DocumentElement] = []
    blocks = pdf_page.get_text("blocks") or []

    for block_index, block in enumerate(blocks):
        text = _pdf_block_text(block)
        if not _looks_like_table_text_block(text):
            continue

        rows = _rows_from_text_block(text)
        metadata = _table_metadata(
            rows=rows,
            detection_method="heuristic",
            page_index_estimated=False,
        )
        metadata.update(
            {
                "source_file": str(path),
                "source_format": "pdf",
                "page_index": page_index,
                "block_index": block_index,
                "table_index": len(elements),
            }
        )
        elements.append(
            DocumentElement(
                element_id=f"pdf-page-{page_number}-table-heuristic-{block_index}",
                element_type="table",
                page_number=page_number,
                text=_flatten_table_text(rows),
                bbox=_bbox_from_value(block[:4], page_number=page_number),
                raw=rows,
                metadata=metadata,
            )
        )

    return elements


def _looks_like_table_text_block(text: str) -> bool:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) < 2:
        return False

    multi_cell_lines = 0
    for line in lines:
        if "\t" in line:
            multi_cell_lines += 1
            continue

        parts = [part for part in line.split("  ") if part.strip()]
        if len(parts) >= 3:
            multi_cell_lines += 1

    return multi_cell_lines >= 2


def _rows_from_text_block(text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if "\t" in line:
            rows.append([cell.strip() for cell in line.split("\t")])
            continue
        rows.append([cell.strip() for cell in re_split_table_columns(line)])
    return rows


def re_split_table_columns(line: str) -> list[str]:
    import re

    parts = [part for part in re.split(r"\s{2,}", line) if part.strip()]
    if len(parts) > 1:
        return parts
    return [line]


def _extract_docx_native_supplement_elements(
    path: Path,
    pagination: str,
) -> list[DocumentElement]:
    doc = Document(str(path))
    elements: list[DocumentElement] = []

    for table_index, table in enumerate(doc.tables):
        elements.append(
            _docx_table_to_element(
                table=table,
                path=path,
                table_index=table_index,
                page_number=1,
                page_index=0,
                pagination=pagination,
                page_index_estimated=True,
            )
        )

    elements.extend(
        _extract_docx_image_elements(
            doc=doc,
            path=path,
            page_number=1,
            page_index=0,
            pagination=pagination,
            page_index_estimated=True,
        )
    )
    return elements


def _docx_table_to_element(
    table: Any,
    path: Path,
    table_index: int,
    page_number: int,
    page_index: int,
    pagination: str,
    page_index_estimated: bool,
) -> DocumentElement:
    rows = _docx_table_to_rows(table)
    metadata = _table_metadata(
        rows=rows,
        detection_method="native_docx",
        page_index_estimated=page_index_estimated,
    )
    metadata.update(
        {
            "source_file": str(path),
            "source_format": "docx",
            "pagination": pagination,
            "page_index": page_index,
            "table_index": table_index,
        }
    )
    return DocumentElement(
        element_id=f"docx-table-{table_index}",
        element_type="table",
        page_number=page_number,
        text=_flatten_table_text(rows),
        raw=rows,
        metadata=metadata,
    )


def _extract_docx_image_elements(
    doc: Any,
    path: Path,
    page_number: int,
    page_index: int,
    pagination: str,
    page_index_estimated: bool,
) -> list[DocumentElement]:
    elements: list[DocumentElement] = []
    for image_index, shape in enumerate(getattr(doc, "inline_shapes", [])):
        elements.append(
            DocumentElement(
                element_id=f"docx-image-{image_index}",
                element_type="image",
                page_number=page_number,
                metadata={
                    "source_file": str(path),
                    "source_format": "docx",
                    "pagination": pagination,
                    "page_index": page_index,
                    "page_index_estimated": page_index_estimated,
                    "image_index": image_index,
                    "width": getattr(shape, "width", None),
                    "height": getattr(shape, "height", None),
                    "extension": None,
                    "detection_method": "native_docx",
                },
            )
        )
    return elements


def _table_metadata(
    rows: list[list[str]],
    detection_method: str,
    page_index_estimated: bool,
) -> dict[str, Any]:
    normalized_rows = _normalize_table_rows(rows)
    row_count = len(normalized_rows)
    col_count = max((len(row) for row in normalized_rows), default=0)
    return {
        "rows": row_count,
        "cols": col_count,
        "cells": normalized_rows,
        "detection_method": detection_method,
        "page_index_estimated": page_index_estimated,
    }


def _normalize_table_rows(rows: Any) -> list[list[str]]:
    normalized: list[list[str]] = []
    for row in rows or []:
        if row is None:
            continue
        normalized.append(["" if cell is None else str(cell).strip() for cell in row])
    return normalized


def _flatten_ocr_table_text(table_item: dict[str, Any]) -> str:
    df = table_item.get("df")
    if df is not None:
        try:
            values = df.fillna("").astype(str).values.flatten()
            parts = [str(value).strip() for value in values if str(value).strip()]
            if parts:
                return " ".join(parts)
        except Exception:
            pass

    spans = table_item.get("spans") or []
    span_text = [
        str(span.get("text", "")).strip()
        for span in spans
        if isinstance(span, dict) and str(span.get("text", "")).strip()
    ]
    if span_text:
        return " ".join(span_text)

    html = table_item.get("html")
    if isinstance(html, str):
        return " ".join(html.split())

    return ""


def _pdf_block_text(block: Any) -> str:
    if isinstance(block, (tuple, list)) and len(block) > 4:
        return str(block[4] or "")
    return ""


def _docx_table_to_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _flatten_table_text(rows: list[list[str]]) -> str:
    parts: list[str] = []
    for row in rows:
        for cell_text in row:
            if cell_text:
                parts.append(cell_text)
    return " ".join(parts)
