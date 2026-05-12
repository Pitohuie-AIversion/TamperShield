from pathlib import Path
from typing import Dict, Iterable, List

import docx
import fitz
import pandas as pd
import pdfplumber

from core.data_normalize import attach_source_metadata, normalize_dataframe


def extract_text_from_pdf_pymupdf(pdf_path: str) -> str:
    texts: List[str] = []

    with fitz.open(pdf_path) as pdf:
        for page in pdf:
            texts.append(page.get_text("text"))

    return "\n".join(texts).strip()


def extract_text_from_pdf_pdfplumber(pdf_path: str) -> str:
    texts: List[str] = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            texts.append(page.extract_text() or "")

    return "\n".join(texts).strip()


def extract_text_from_docx(docx_path: str) -> str:
    document = docx.Document(docx_path)
    return "\n".join(p.text for p in document.paragraphs).strip()


def _extract_docx_table_rows(table: docx.table.Table) -> List[List[str]]:
    rows: List[List[str]] = []

    for row in table.rows:
        rows.append(
            [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]
        )

    return rows


def detect_file_type(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        return "pdf"

    if suffix == ".docx":
        return "docx"

    raise ValueError(f"Unsupported file type: {suffix}")


def is_supported_native_document(file_path: str) -> bool:
    path = Path(file_path)

    if path.name.startswith("~$"):
        return False

    return path.suffix.lower() in {".pdf", ".docx"}


def extract_tables_from_native_pdf_with_metadata(pdf_path: str) -> List[Dict[str, object]]:
    tables: List[Dict[str, object]] = []

    with pdfplumber.open(pdf_path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            raw_tables = page.extract_tables()

            for table_index, table in enumerate(raw_tables):
                if not table:
                    continue

                cleaned_table = [
                    [cell if cell is not None else "" for cell in row]
                    for row in table
                ]

                if len(cleaned_table) <= 1:
                    continue

                header = cleaned_table[0]
                rows = cleaned_table[1:]

                df = pd.DataFrame(rows, columns=header)
                df = normalize_dataframe(df)
                df = attach_source_metadata(
                    df,
                    source_file=pdf_path,
                    page_index=page_index,
                    table_index=table_index,
                )

                tables.append(
                    {
                        "df": df,
                        "source_file": pdf_path,
                        "page_index": page_index,
                        "table_index": table_index,
                    }
                )

    return tables


def extract_tables_from_native_pdf(pdf_path: str) -> List[pd.DataFrame]:
    tables = extract_tables_from_native_pdf_with_metadata(pdf_path)
    return [item["df"] for item in tables]


def extract_tables_from_native_docx_with_metadata(
    docx_path: str,
) -> List[Dict[str, object]]:
    tables: List[Dict[str, object]] = []
    document = docx.Document(docx_path)

    for table_index, table in enumerate(document.tables):
        raw_rows = _extract_docx_table_rows(table)

        if len(raw_rows) <= 1:
            continue

        header = raw_rows[0]
        rows = raw_rows[1:]

        df = pd.DataFrame(rows, columns=header)
        df = normalize_dataframe(df)
        df = attach_source_metadata(
            df,
            source_file=docx_path,
            table_index=table_index,
        )

        tables.append(
            {
                "df": df,
                "source_file": docx_path,
                "table_index": table_index,
            }
        )

    return tables


def extract_tables_from_native_docx(docx_path: str) -> List[pd.DataFrame]:
    tables = extract_tables_from_native_docx_with_metadata(docx_path)
    return [item["df"] for item in tables]


def extract_tables_from_native_document_with_metadata(
    file_path: str,
) -> List[Dict[str, object]]:
    file_type = detect_file_type(file_path)

    if file_type == "pdf":
        return extract_tables_from_native_pdf_with_metadata(file_path)

    if file_type == "docx":
        return extract_tables_from_native_docx_with_metadata(file_path)

    raise ValueError(f"Unsupported file type: {file_type}")


def extract_tables_from_native_document(file_path: str) -> List[pd.DataFrame]:
    tables = extract_tables_from_native_document_with_metadata(file_path)
    return [item["df"] for item in tables]


def extract_tables_from_native_documents_with_metadata(
    file_paths: Iterable[str],
) -> List[Dict[str, object]]:
    tables: List[Dict[str, object]] = []

    for file_path in file_paths:
        if not is_supported_native_document(file_path):
            continue

        tables.extend(extract_tables_from_native_document_with_metadata(file_path))

    return tables


def extract_tables_from_native_documents(
    file_paths: Iterable[str],
) -> List[pd.DataFrame]:
    tables = extract_tables_from_native_documents_with_metadata(file_paths)
    return [item["df"] for item in tables]
